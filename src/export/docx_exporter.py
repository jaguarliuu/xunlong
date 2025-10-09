"""
DOCX - /Word
"""
from pathlib import Path
from typing import Dict, Any, Optional
from loguru import logger


class DOCXExporter:
    """DOCX"""

    async def export(
        self,
        project_dir: Path,
        output_path: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        DOCX

        Args:
            project_dir: 
            output_path: 

        Returns:
            
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            # Markdown
            md_file = project_dir / "reports" / "FINAL_REPORT.md"
            if not md_file.exists():
                return {
                    "status": "error",
                    "error": "Markdown"
                }

            # 
            with open(md_file, 'r', encoding='utf-8') as f:
                md_content = f.read()

            logger.info(f"DOCX: {md_file}")

            # Word
            doc = Document()

            # Markdown
            self._parse_markdown_to_docx(doc, md_content)

            # 
            if not output_path:
                output_path = project_dir / "exports" / "report.docx"
            else:
                output_path = Path(output_path)

            output_path.parent.mkdir(parents=True, exist_ok=True)

            # 
            doc.save(str(output_path))

            # 
            file_size = output_path.stat().st_size
            file_size_str = self._format_file_size(file_size)

            logger.info(f"DOCX: {output_path}")

            return {
                "status": "success",
                "output_file": str(output_path),
                "file_size": file_size_str
            }

        except ImportError:
            return {
                "status": "error",
                "error": "python-docx: pip install python-docx"
            }
        except Exception as e:
            logger.error(f"DOCX: {e}")
            import traceback
            traceback.print_exc()
            return {
                "status": "error",
                "error": str(e)
            }

    def _parse_markdown_to_docx(self, doc, md_content: str):
        """
        MarkdownDOCX
        """
        from docx.shared import Pt

        lines = md_content.split('\n')

        for line in lines:
            line = line.strip()

            if not line:
                continue

            # 
            if line.startswith('# '):
                p = doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                p = doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                p = doc.add_heading(line[4:], level=3)

            # 
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')

            # 
            elif line[0].isdigit() and '. ' in line[:4]:
                parts = line.split('. ', 1)
                if len(parts) == 2:
                    p = doc.add_paragraph(parts[1], style='List Number')

            # 
            else:
                p = doc.add_paragraph(line)
                p.paragraph_format.line_spacing = 1.5

    def _format_file_size(self, size_bytes: int) -> str:
        """TODO: Add docstring."""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.2f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.2f} TB"
