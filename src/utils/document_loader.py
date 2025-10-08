"""
Document loader utilities for ingesting user-provided context files.

Supports plain text (.txt), PDF (.pdf), and Word documents (.docx).
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Tuple

from loguru import logger


class DocumentLoadError(Exception):
    """Raised when a document cannot be loaded or parsed."""


@dataclass
class LoadedDocument:
    """Structured representation of a loaded document."""

    content: str
    filename: str
    suffix: str
    char_length: int
    truncated: bool
    source_path: str


MAX_DOCUMENT_CHARACTERS = 20_000


def load_document(path: str | Path) -> LoadedDocument:
    """
    Load a document from the given filesystem path.

    Args:
        path: Path to the document file.

    Returns:
        LoadedDocument: Parsed document content and metadata.

    Raises:
        DocumentLoadError: If the file does not exist or cannot be parsed.
    """
    file_path = Path(path)
    if not file_path.exists():
        raise DocumentLoadError(f"Document not found: {file_path}")

    suffix = file_path.suffix.lower()
    try:
        if suffix == ".txt":
            content = _load_txt(file_path)
        elif suffix in {".docx"}:
            content = _load_docx(file_path)
        elif suffix == ".pdf":
            content = _load_pdf(file_path)
        else:
            raise DocumentLoadError(
                f"Unsupported document format '{suffix}'. "
                "Supported formats are: .txt, .pdf, .docx"
            )
    except DocumentLoadError:
        raise
    except Exception as exc:  # pragma: no cover - unforeseen parsing errors
        raise DocumentLoadError(f"Failed to parse document: {exc}") from exc

    if not content:
        raise DocumentLoadError("Document appears to be empty after parsing")

    normalized = content.replace("\r\n", "\n").replace("\r", "\n").strip()
    if not normalized:
        raise DocumentLoadError("Document contains no readable text")

    truncated = False
    if len(normalized) > MAX_DOCUMENT_CHARACTERS:
        logger.warning(
            "Document content exceeds %d characters; truncating for prompt safety",
            MAX_DOCUMENT_CHARACTERS,
        )
        normalized = normalized[:MAX_DOCUMENT_CHARACTERS]
        truncated = True

    return LoadedDocument(
        content=normalized,
        filename=file_path.name,
        suffix=suffix,
        char_length=len(normalized),
        truncated=truncated,
        source_path=str(file_path.resolve()),
    )


def _load_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _load_docx(path: Path) -> str:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - dependency missing
        raise DocumentLoadError(
            "python-docx is required to parse .docx files"
        ) from exc

    document = docx.Document(str(path))
    parts = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    # Extract text from tables if any
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _load_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency missing
        raise DocumentLoadError("pypdf is required to parse .pdf files") from exc

    reader = PdfReader(str(path))
    pages_text = []
    for page_number, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:  # pragma: no cover - pdf quirks
            logger.warning("Failed to extract text from page %d: %s", page_number, exc)
            text = ""
        if text.strip():
            pages_text.append(text.strip())

    return "\n\n".join(pages_text)
